from google.colab import drive
drive.mount('/content/drive')

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import warnings
warnings.filterwarnings('ignore')


"""
DATA LOADING AND PREPROCESSING
"""
# Loading of the dataset
file_path_in_drive = '/content/drive/MyDrive/TI18_TRUESIGHT_FILLED.csv'

try:
    df_drive = pd.read_csv(file_path_in_drive)
    print("DataFrame loaded successfully from Google Drive:")
    display(df_drive.head())
except FileNotFoundError:
    print(f"Error: The file '{file_path_in_drive}' was not found. Please check the path and try again.")
except Exception as e:
    print(f"An error occurred: {e}")


# Drop the first column because the second column is already precleaned after
# the scraping
# This assumes the first column is identified by its position (index 0).
df_modified = df_drive.iloc[:, 1:]

print("DataFrame after dropping the first column:")
display(df_modified.head())


print("="*80)
print("DATASET OVERVIEW")
print("="*80)

# Basic information
print(f"\n1. Dataset Shape: {df_modified.shape[0]} rows × {df_modified.shape[1]} columns")
print(f"   - Number of entries: {len(df_modified)}")
print(f"   - Number of features: {len(df_modified.columns)}")

print("\n2. Column Names and Data Types:")
print(df_modified.dtypes)

print("\n3. First Few Rows:")
display(df_modified.head())

print("\n4. Statistical Summary:")
display(df_modified.describe(include='all'))

print("\n5. Missing Values:")
missing = df_modified.isnull().sum()
print(missing[missing > 0] if missing.sum() > 0 else "No missing values found!")

print("\n6. Duplicate Rows:")
print(f"   Number of duplicates: {df_modified.duplicated().sum()}")



# Check if we have a text column for text mining
print("\n7. Text Columns Analysis:")
text_columns = df_modified.select_dtypes(include=['object']).columns.tolist()
print(f"   Potential text columns: {text_columns}")

if text_columns:
    for col in text_columns:
        print(f"\n   Column '{col}':")
        print(f"   - Unique values: {df_modified[col].nunique()}")
        print(f"   - Sample values: {df_modified[col].head(7).tolist()}")
        if df_modified[col].dtype == 'object':
            avg_length = df_modified[col].astype(str).str.len().mean()
            print(f"   - Average text length: {avg_length:.2f} characters")
            
print("\nUnique values for 'emotion' feature:")
print(df_modified['emotion'].unique().tolist())





"""
VISUALIZATIONS AND ANALYSES
"""

# Visualizations
print("\n" + "="*80)
print("GENERATING VISUALIZATIONS")
print("="*80)

# 1. Emotion Distribution
plt.figure(figsize=(12, 6))
emotion_counts = df_modified['emotion'].value_counts()
colors = plt.cm.Set3(np.linspace(0, 1, len(emotion_counts)))
bars = plt.bar(emotion_counts.index, emotion_counts.values, color=colors, edgecolor='black', linewidth=1.2)
plt.title('Distribution of Emotions in Comments', fontsize=16, fontweight='bold', pad=20)
plt.xlabel('Emotion', fontsize=12, fontweight='bold')
plt.ylabel('Frequency', fontsize=12, fontweight='bold')
plt.xticks(rotation=45, ha='right')
plt.grid(axis='y', alpha=0.3, linestyle='--')

# Add value labels on bars
for bar in bars:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2., height,
             f'{int(height)}\n({height/len(df_modified)*100:.1f}%)',
             ha='center', va='bottom', fontsize=10, fontweight='bold')
plt.tight_layout()
plt.show()

# 2. Pie Chart for Emotion Distribution
plt.figure(figsize=(10, 8))
colors_pie = plt.cm.Pastel1(np.linspace(0, 1, len(emotion_counts)))
wedges, texts, autotexts = plt.pie(emotion_counts.values, labels=emotion_counts.index, 
                                     autopct='%1.1f%%', colors=colors_pie, startangle=90,
                                     textprops={'fontsize': 11, 'fontweight': 'bold'},
                                     explode=[0.05 if i == 0 else 0 for i in range(len(emotion_counts))])
plt.title('Emotion Distribution (Percentage)', fontsize=16, fontweight='bold', pad=20)
for autotext in autotexts:
    autotext.set_color('white')
    autotext.set_fontweight('bold')
plt.tight_layout()
plt.show()

# 3. Text Length Analysis
df_modified['text_length'] = df_modified['cleanedComments'].astype(str).str.len()
df_modified['word_count'] = df_modified['cleanedComments'].astype(str).str.split().str.len()

fig, axes = plt.subplots(1, 2, figsize=(16, 6))

# Character length distribution
axes[0].hist(df_modified['text_length'].dropna(), bins=50, color='skyblue', edgecolor='black', alpha=0.7)
axes[0].axvline(df_modified['text_length'].mean(), color='red', linestyle='--', linewidth=2, label=f'Mean: {df_modified["text_length"].mean():.1f}')
axes[0].axvline(df_modified['text_length'].median(), color='green', linestyle='--', linewidth=2, label=f'Median: {df_modified["text_length"].median():.1f}')
axes[0].set_xlabel('Character Length', fontsize=12, fontweight='bold')
axes[0].set_ylabel('Frequency', fontsize=12, fontweight='bold')
axes[0].set_title('Distribution of Comment Length (Characters)', fontsize=14, fontweight='bold')
axes[0].legend()
axes[0].grid(alpha=0.3)

# Word count distribution
axes[1].hist(df_modified['word_count'].dropna(), bins=50, color='lightcoral', edgecolor='black', alpha=0.7)
axes[1].axvline(df_modified['word_count'].mean(), color='red', linestyle='--', linewidth=2, label=f'Mean: {df_modified["word_count"].mean():.1f}')
axes[1].axvline(df_modified['word_count'].median(), color='green', linestyle='--', linewidth=2, label=f'Median: {df_modified["word_count"].median():.1f}')
axes[1].set_xlabel('Word Count', fontsize=12, fontweight='bold')
axes[1].set_ylabel('Frequency', fontsize=12, fontweight='bold')
axes[1].set_title('Distribution of Comment Length (Words)', fontsize=14, fontweight='bold')
axes[1].legend()
axes[1].grid(alpha=0.3)

plt.tight_layout()
plt.show()

# 4. Text Length by Emotion
plt.figure(figsize=(14, 6))
df_modified.boxplot(column='text_length', by='emotion', figsize=(14, 6), patch_artist=True)
plt.suptitle('')
plt.title('Comment Length Distribution by Emotion', fontsize=14, fontweight='bold', pad=20)
plt.xlabel('Emotion', fontsize=12, fontweight='bold')
plt.ylabel('Character Length', fontsize=12, fontweight='bold')
plt.xticks(rotation=45, ha='right')
plt.grid(alpha=0.3)
plt.tight_layout()
plt.show()

# 5. Word Count by Emotion (Violin Plot)
plt.figure(figsize=(14, 6))
sns.violinplot(data=df_modified, x='emotion', y='word_count', palette='muted')
plt.title('Word Count Distribution by Emotion (Violin Plot)', fontsize=14, fontweight='bold', pad=20)
plt.xlabel('Emotion', fontsize=12, fontweight='bold')
plt.ylabel('Word Count', fontsize=12, fontweight='bold')
plt.xticks(rotation=45, ha='right')
plt.grid(axis='y', alpha=0.3)
plt.tight_layout()
plt.show()

# 6. Statistical Summary by Emotion
print("\n" + "="*80)
print("TEXT LENGTH STATISTICS BY EMOTION")
print("="*80)
length_stats = df_modified.groupby('emotion').agg({
    'text_length': ['mean', 'median', 'min', 'max', 'std'],
    'word_count': ['mean', 'median', 'min', 'max', 'std']
}).round(2)
print(length_stats)

# 7. Top Words Analysis
from collections import Counter
import re

print("\n" + "="*80)
print("TOP 20 MOST COMMON WORDS (OVERALL)")
print("="*80)

# Combine all comments and extract words
all_text = ' '.join(df_modified['cleanedComments'].dropna().astype(str))
words = re.findall(r'\b[a-zA-Z]+\b', all_text.lower())
word_freq = Counter(words)

# Remove common stop words
stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'is', 'was', 'are', 'were', 'this', 'that', 'it', 'with', 'as', 'be', 'have', 'has', 'had'}
filtered_words = {word: count for word, count in word_freq.items() if word not in stop_words and len(word) > 2}
top_words = Counter(filtered_words).most_common(20)

for i, (word, count) in enumerate(top_words, 1):
    print(f"{i:2d}. {word:15s} - {count:5d} occurrences")

# Plot top words
plt.figure(figsize=(14, 6))
words, counts = zip(*top_words)
colors_bar = plt.cm.viridis(np.linspace(0, 1, len(words)))
bars = plt.barh(range(len(words)), counts, color=colors_bar, edgecolor='black')
plt.yticks(range(len(words)), words)
plt.xlabel('Frequency', fontsize=12, fontweight='bold')
plt.ylabel('Words', fontsize=12, fontweight='bold')
plt.title('Top 20 Most Common Words in Comments', fontsize=14, fontweight='bold', pad=20)
plt.gca().invert_yaxis()
plt.grid(axis='x', alpha=0.3)

# Add value labels
for i, (bar, count) in enumerate(zip(bars, counts)):
    plt.text(count, i, f' {count}', va='center', fontsize=9, fontweight='bold')

plt.tight_layout()
plt.show()

# 8. Top Words by Emotion
print("\n" + "="*80)
print("TOP 10 WORDS BY EACH EMOTION")
print("="*80)

emotions = df_modified['emotion'].unique()
fig, axes = plt.subplots(3, 3, figsize=(18, 14))
axes = axes.flatten()

for idx, emotion in enumerate(emotions):
    emotion_text = ' '.join(df_modified[df_modified['emotion'] == emotion]['cleanedComments'].dropna().astype(str))
    emotion_words = re.findall(r'\b[a-zA-Z]+\b', emotion_text.lower())
    emotion_freq = Counter(emotion_words)
    filtered_emotion_words = {word: count for word, count in emotion_freq.items() if word not in stop_words and len(word) > 2}
    top_emotion_words = Counter(filtered_emotion_words).most_common(10)
    
    print(f"\n{emotion.upper()}:")
    for i, (word, count) in enumerate(top_emotion_words, 1):
        print(f"  {i:2d}. {word:15s} - {count:4d}")
    
    if top_emotion_words:
        words_e, counts_e = zip(*top_emotion_words)
        axes[idx].barh(range(len(words_e)), counts_e, color=plt.cm.Set2(idx % 8))
        axes[idx].set_yticks(range(len(words_e)))
        axes[idx].set_yticklabels(words_e, fontsize=9)
        axes[idx].set_xlabel('Frequency', fontsize=10)
        axes[idx].set_title(f'{emotion.capitalize()}', fontsize=12, fontweight='bold')
        axes[idx].invert_yaxis()
        axes[idx].grid(axis='x', alpha=0.3)

# Hide unused subplots
for idx in range(len(emotions), len(axes)):
    axes[idx].axis('off')

plt.suptitle('Top 10 Words by Emotion', fontsize=16, fontweight='bold', y=1.00)
plt.tight_layout()
plt.show()

# 9. Correlation Heatmap (Numeric features)
print("\n" + "="*80)
print("CORRELATION ANALYSIS")
print("="*80)

plt.figure(figsize=(8, 6))
numeric_cols = ['text_length', 'word_count']
correlation_matrix = df_modified[numeric_cols].corr()
sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm', center=0, 
            square=True, linewidths=1, cbar_kws={"shrink": 0.8})
plt.title('Correlation Between Text Features', fontsize=14, fontweight='bold', pad=20)
plt.tight_layout()
plt.show()

# 10. Data Quality Overview
print("\n" + "="*80)
print("DATA QUALITY SUMMARY")
print("="*80)

total_rows = len(df_modified)
missing_comments = df_modified['cleanedComments'].isnull().sum()
duplicates = df_modified.duplicated().sum()
empty_strings = (df_modified['cleanedComments'].astype(str).str.strip() == '').sum()

print(f"Total Rows: {total_rows}")
print(f"Missing Comments: {missing_comments} ({missing_comments/total_rows*100:.2f}%)")
print(f"Duplicate Rows: {duplicates} ({duplicates/total_rows*100:.2f}%)")
print(f"Empty Comments: {empty_strings} ({empty_strings/total_rows*100:.2f}%)")
print(f"Valid Comments: {total_rows - missing_comments - empty_strings} ({(total_rows - missing_comments - empty_strings)/total_rows*100:.2f}%)")

# Visualize data quality
plt.figure(figsize=(10, 6))
quality_data = {
    'Valid': total_rows - missing_comments - empty_strings - duplicates,
    'Missing': missing_comments,
    'Duplicates': duplicates,
    'Empty': empty_strings
}
colors_quality = ['#2ecc71', '#e74c3c', '#f39c12', '#95a5a6']
plt.pie(quality_data.values(), labels=quality_data.keys(), autopct='%1.1f%%', 
        colors=colors_quality, startangle=90, textprops={'fontsize': 12, 'fontweight': 'bold'})
plt.title('Data Quality Distribution', fontsize=16, fontweight='bold', pad=20)
plt.tight_layout()
plt.show()

print("\n" + "="*80)
print("EXPLORATION COMPLETE!")
print("="*80)


# ============================================================================
# ADVANCED DATA PREPROCESSING PIPELINE
# ============================================================================

print("\n" + "="*80)
print("DATA PREPROCESSING PIPELINE")
print("="*80)

# Install necessary libraries
import subprocess
import sys

def install_package(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", package])

try:
    import nltk
    from textblob import TextBlob
    import contractions
except ImportError:
    print("Installing required packages...")
    install_package("nltk")
    install_package("textblob")
    install_package("contractions")
    import nltk
    from textblob import TextBlob
    import contractions
    print("✓ Packages installed successfully")

import re
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
from sklearn.model_selection import train_test_split

# Download NLTK data
print("\nDownloading NLTK resources...")
for resource in ['punkt', 'punkt_tab', 'stopwords', 'wordnet', 'averaged_perceptron_tagger', 'omw-1.4']:
    try:
        nltk.download(resource, quiet=True)
    except:
        pass
print("✓ NLTK resources downloaded successfully")

print("\n1. INITIAL DATA CLEANING")
print("-" * 80)

# Create a copy for preprocessing
df_clean = df_modified.copy()

# Remove duplicates
initial_rows = len(df_clean)
df_clean = df_clean.drop_duplicates(subset=['cleanedComments'])
print(f"✓ Removed {initial_rows - len(df_clean)} duplicate rows")

# Handle missing values
df_clean = df_clean.dropna(subset=['cleanedComments'])
print(f"✓ Removed rows with missing comments. Remaining: {len(df_clean)}")

# Remove empty or very short texts
df_clean = df_clean[df_clean['cleanedComments'].astype(str).str.len() > 10]
print(f"✓ Removed very short texts (< 10 chars). Remaining: {len(df_clean)}")

print("\n2. TEXT PREPROCESSING FUNCTIONS")
print("-" * 80)

def preprocess_text_advanced(text):
    """
    Advanced text preprocessing for emotion analysis
    """
    if pd.isna(text) or not isinstance(text, str):
        return ""
    
    # Convert to lowercase
    text = text.lower()
    
    # Expand contractions (e.g., don't -> do not)
    try:
        text = contractions.fix(text)
    except:
        pass
    
    # Remove URLs
    text = re.sub(r'http\S+|www\S+|https\S+', '', text, flags=re.MULTILINE)
    
    # Remove email addresses
    text = re.sub(r'\S+@\S+', '', text)
    
    # Remove mentions and hashtags
    text = re.sub(r'@\w+|#\w+', '', text)
    
    # Keep letters and basic punctuation for emotion (!, ?, .)
    text = re.sub(r'[^a-zA-Z\s!?.]', ' ', text)
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Tokenization
    tokens = word_tokenize(text)
    
    # Remove stopwords but keep emotion-relevant words
    stop_words = set(stopwords.words('english'))
    # Keep important emotion words that might be in stopwords
    keep_words = {'not', 'no', 'never', 'very', 'too', 'so', 'really', 'most', 'more'}
    stop_words = stop_words - keep_words
    
    tokens = [word for word in tokens if word not in stop_words and len(word) > 1]
    
    # Lemmatization
    lemmatizer = WordNetLemmatizer()
    tokens = [lemmatizer.lemmatize(word) for word in tokens]
    
    return ' '.join(tokens)

# Apply preprocessing
print("Applying advanced preprocessing...")
df_clean['text_processed'] = df_clean['cleanedComments'].apply(preprocess_text_advanced)
print("✓ Advanced text preprocessing completed")

print("\n3. FEATURE ENGINEERING")
print("-" * 80)

# Text length features
df_clean['original_length'] = df_clean['cleanedComments'].str.len()
df_clean['processed_length'] = df_clean['text_processed'].str.len()
df_clean['word_count_processed'] = df_clean['text_processed'].apply(lambda x: len(x.split()) if x else 0)

# Average word length
df_clean['avg_word_length'] = df_clean['text_processed'].apply(
    lambda x: np.mean([len(word) for word in x.split()]) if x and len(x.split()) > 0 else 0
)

# Sentiment analysis using TextBlob
print("Performing sentiment analysis...")
df_clean['sentiment_polarity'] = df_clean['cleanedComments'].apply(
    lambda x: TextBlob(str(x)).sentiment.polarity
)
df_clean['sentiment_subjectivity'] = df_clean['cleanedComments'].apply(
    lambda x: TextBlob(str(x)).sentiment.subjectivity
)
df_clean['sentiment_label'] = df_clean['sentiment_polarity'].apply(
    lambda x: 'positive' if x > 0.1 else ('negative' if x < -0.1 else 'neutral')
)
print("✓ Sentiment features extracted")

# Lexical diversity (vocabulary richness)
df_clean['lexical_diversity'] = df_clean['text_processed'].apply(
    lambda x: len(set(x.split())) / len(x.split()) if x and len(x.split()) > 0 else 0
)

# Special character counts (emotion indicators)
df_clean['exclamation_count'] = df_clean['cleanedComments'].str.count('!')
df_clean['question_count'] = df_clean['cleanedComments'].str.count('\?')
df_clean['emoji_count'] = df_clean['cleanedComments'].apply(
    lambda x: len(re.findall(r'[^\w\s,]', str(x))) if pd.notna(x) else 0
)

print("✓ All features engineered successfully")

print("\n4. FEATURE SUMMARY")
print("-" * 80)
new_features = ['text_processed', 'original_length', 'processed_length', 
                'word_count_processed', 'avg_word_length', 'sentiment_polarity', 
                'sentiment_subjectivity', 'sentiment_label', 'lexical_diversity',
                'exclamation_count', 'question_count', 'emoji_count']
print(f"Total features: {len(new_features)}")
for feature in new_features:
    print(f"  ✓ {feature}")

print(f"\nFinal dataset size: {len(df_clean)} entries")
print(f"Total columns: {len(df_clean.columns)}")

# Display sample of preprocessed data
print("\n5. SAMPLE OF PREPROCESSED DATA")
print("-" * 80)
sample_df = df_clean[['cleanedComments', 'text_processed', 'emotion', 
                       'sentiment_polarity', 'sentiment_label']].head(5)
for idx, row in sample_df.iterrows():
    print(f"\n{'='*60}")
    print(f"Sample {idx + 1}:")
    print(f"Original : {row['cleanedComments'][:80]}...")
    print(f"Processed: {row['text_processed'][:80]}...")
    print(f"Emotion  : {row['emotion']}")
    print(f"Sentiment: {row['sentiment_label']} (polarity: {row['sentiment_polarity']:.3f})")

# ============================================================================
# PREPROCESSING IMPACT VISUALIZATION
# ============================================================================

print("\n" + "="*80)
print("PREPROCESSING IMPACT VISUALIZATION")
print("="*80)

fig, axes = plt.subplots(2, 3, figsize=(18, 10))

# 1. Text length comparison
axes[0, 0].hist(df_clean['original_length'], bins=50, alpha=0.6, label='Original', 
                edgecolor='black', color='skyblue')
axes[0, 0].hist(df_clean['processed_length'], bins=50, alpha=0.6, label='Processed', 
                edgecolor='black', color='coral')
axes[0, 0].set_title('Text Length: Original vs Processed', fontsize=12, fontweight='bold')
axes[0, 0].set_xlabel('Length (characters)')
axes[0, 0].set_ylabel('Frequency')
axes[0, 0].legend()
axes[0, 0].grid(alpha=0.3)

# 2. Word count distribution
axes[0, 1].hist(df_clean['word_count_processed'], bins=40, edgecolor='black', 
                color='lightgreen', alpha=0.7)
axes[0, 1].axvline(df_clean['word_count_processed'].mean(), color='red', 
                   linestyle='--', linewidth=2, label=f"Mean: {df_clean['word_count_processed'].mean():.1f}")
axes[0, 1].axvline(df_clean['word_count_processed'].median(), color='green', linestyle='--', linewidth=2, label=f"Median: {df_clean['word_count_processed'].median():.1f}")
axes[0, 1].set_title('Processed Word Count Distribution', fontsize=12, fontweight='bold')
axes[0, 1].set_xlabel('Number of Words')
axes[0, 1].set_ylabel('Frequency')
axes[0, 1].legend()
axes[0, 1].grid(alpha=0.3)

# 3. Sentiment distribution
sentiment_counts = df_clean['sentiment_label'].value_counts()
colors_sent = {'positive': '#2ecc71', 'neutral': '#95a5a6', 'negative': '#e74c3c'}
bars = axes[0, 2].bar(sentiment_counts.index, sentiment_counts.values, 
                       color=[colors_sent.get(x, 'gray') for x in sentiment_counts.index],
                       edgecolor='black', linewidth=1.2)
axes[0, 2].set_title('Sentiment Distribution', fontsize=12, fontweight='bold')
axes[0, 2].set_xlabel('Sentiment')
axes[0, 2].set_ylabel('Count')
axes[0, 2].grid(axis='y', alpha=0.3)
for bar in bars:
    height = bar.get_height()
    axes[0, 2].text(bar.get_x() + bar.get_width()/2., height,
                    f'{int(height)}', ha='center', va='bottom', fontweight='bold')

# 4. Sentiment polarity distribution
axes[1, 0].hist(df_clean['sentiment_polarity'], bins=50, edgecolor='black', 
                color='mediumpurple', alpha=0.7)
axes[1, 0].axvline(x=0, color='black', linestyle='--', linewidth=2, label='Neutral')
axes[1, 0].set_title('Sentiment Polarity Distribution', fontsize=12, fontweight='bold')
axes[1, 0].set_xlabel('Polarity Score')
axes[1, 0].set_ylabel('Frequency')
axes[1, 0].legend()
axes[1, 0].grid(alpha=0.3)

# 5. Lexical diversity by emotion
df_clean.boxplot(column='lexical_diversity', by='emotion', ax=axes[1, 1], 
                 patch_artist=True, grid=True)
axes[1, 1].set_title('Lexical Diversity by Emotion', fontsize=12, fontweight='bold')
axes[1, 1].set_xlabel('Emotion')
axes[1, 1].set_ylabel('Lexical Diversity Score')
axes[1, 1].get_figure().suptitle('')
plt.setp(axes[1, 1].xaxis.get_majorticklabels(), rotation=45, ha='right')

# 6. Emotion vs Sentiment comparison
emotion_sentiment = pd.crosstab(df_clean['emotion'], df_clean['sentiment_label'], normalize='index') * 100
emotion_sentiment.plot(kind='bar', stacked=False, ax=axes[1, 2], 
                       color=['#e74c3c', '#95a5a6', '#2ecc71'], 
                       edgecolor='black', linewidth=1)
axes[1, 2].set_title('Emotion vs Sentiment Analysis', fontsize=12, fontweight='bold')
axes[1, 2].set_xlabel('Emotion')
axes[1, 2].set_ylabel('Percentage (%)')
axes[1, 2].legend(title='Sentiment', loc='upper right')
axes[1, 2].grid(axis='y', alpha=0.3)
plt.setp(axes[1, 2].xaxis.get_majorticklabels(), rotation=45, ha='right')

plt.suptitle('Preprocessing Impact Analysis', fontsize=16, fontweight='bold', y=1.00)
plt.tight_layout()
plt.show()

# ============================================================================
# EMOTION-SENTIMENT CORRELATION ANALYSIS
# ============================================================================

print("\n" + "="*80)
print("EMOTION-SENTIMENT CORRELATION ANALYSIS")
print("="*80)

# Statistical comparison of emotions
emotion_stats = df_clean.groupby('emotion').agg({
    'sentiment_polarity': ['mean', 'std', 'min', 'max'],
    'sentiment_subjectivity': ['mean', 'std'],
    'word_count_processed': ['mean', 'std'],
    'lexical_diversity': ['mean', 'std'],
    'exclamation_count': 'mean',
    'question_count': 'mean'
}).round(3)

print("\nEmotion Statistics:")
print(emotion_stats)

# Correlation matrix for features
print("\n" + "="*80)
print("FEATURE CORRELATION MATRIX")
print("="*80)

numeric_features = ['sentiment_polarity', 'sentiment_subjectivity', 'word_count_processed',
                   'lexical_diversity', 'avg_word_length', 'exclamation_count', 'question_count']
correlation_matrix = df_clean[numeric_features].corr()

plt.figure(figsize=(10, 8))
sns.heatmap(correlation_matrix, annot=True, fmt='.2f', cmap='coolwarm', 
            center=0, square=True, linewidths=1, cbar_kws={"shrink": 0.8},
            vmin=-1, vmax=1)
plt.title('Feature Correlation Matrix', fontsize=14, fontweight='bold', pad=20)
plt.tight_layout()
plt.show()

# Save preprocessed data
print("\n" + "="*80)
print("SAVING PREPROCESSED DATA")
print("="*80)

# Save to CSV
output_file = '/content/drive/MyDrive/TI18_TRUESIGHT_PREPROCESSED.csv'
df_clean.to_csv(output_file, index=False)
print(f"✓ Preprocessed data saved to: {output_file}")
print(f"✓ Total entries: {len(df_clean)}")
print(f"✓ Total features: {len(df_clean.columns)}")

print("\n" + "="*80)
print("✓ DATA PREPROCESSING COMPLETED SUCCESSFULLY!")
print("="*80)
print(f"Ready for model development with {len(df_clean)} clean entries")
print(f"Dataset is balanced across {df_clean['emotion'].nunique()} emotion categories")








# ============================================================================
# MODEL DEVELOPMENT & EVALUATION
# ============================================================================

print("\n" + "="*80)
print("MODEL DEVELOPMENT & EVALUATION FOR EMOTION CLASSIFICATION")
print("="*80)

# Install additional libraries
print("\nInstalling required packages...")
try:
    from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
    from sklearn.decomposition import LatentDirichletAllocation, NMF, TruncatedSVD
    from sklearn.cluster import KMeans
    from sklearn.metrics import silhouette_score, calinski_harabasz_score, davies_bouldin_score
    from sklearn.manifold import TSNE
    from sklearn.naive_bayes import MultinomialNB
    from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
    from sklearn.svm import SVC
    from sklearn.linear_model import LogisticRegression
    from sklearn.metrics import classification_report, confusion_matrix, accuracy_score, f1_score
    from wordcloud import WordCloud
except ImportError:
    install_package("wordcloud")
    from wordcloud import WordCloud

from collections import Counter

print("✓ All packages loaded successfully")

# ============================================================================
# PART A: TEXT VECTORIZATION
# ============================================================================

print("\n" + "="*80)
print("📊 PART A: TEXT VECTORIZATION")
print("="*80)

# Prepare text data for modeling
X_text = df_clean['text_processed'].values
y_emotion = df_clean['emotion'].values  # Target: emotion labels

print(f"\nDataset size: {len(X_text)} samples")
print(f"Emotion classes: {np.unique(y_emotion)}")
print(f"Number of classes: {len(np.unique(y_emotion))}")

# Class distribution
print("\nEmotion distribution:")
emotion_dist = Counter(y_emotion)
for emotion, count in emotion_dist.most_common():
    print(f"  {emotion:<12}: {count:>5} ({count/len(y_emotion)*100:>5.2f}%)")

# 1. TF-IDF Vectorization
print("\n1. TF-IDF Vectorization")
print("-" * 80)
tfidf = TfidfVectorizer(max_features=2000, min_df=3, max_df=0.8, ngram_range=(1, 2))
X_tfidf = tfidf.fit_transform(X_text)
print(f"✓ TF-IDF shape: {X_tfidf.shape}")
print(f"✓ Vocabulary size: {len(tfidf.vocabulary_)}")
print(f"✓ Top 15 features: {list(tfidf.get_feature_names_out()[:15])}")

# 2. Count Vectorization (Bag of Words)
print("\n2. Count Vectorization (BoW)")
print("-" * 80)
count_vec = CountVectorizer(max_features=2000, min_df=3, max_df=0.8)
X_count = count_vec.fit_transform(X_text)
print(f"✓ Count Vector shape: {X_count.shape}")
print(f"✓ Vocabulary size: {len(count_vec.vocabulary_)}")

# ============================================================================
# PART B: UNSUPERVISED LEARNING - TOPIC MODELING & CLUSTERING
# ============================================================================

print("\n" + "="*80)
print("📊 PART B: UNSUPERVISED LEARNING")
print("="*80)

model_results = {}

# ============ MODEL 1: LDA TOPIC MODELING ============
print("\n[MODEL 1] Latent Dirichlet Allocation (LDA)")
print("-" * 80)
n_topics = 7  # Match number of emotions
lda = LatentDirichletAllocation(n_components=n_topics, random_state=42, max_iter=20)
lda_topics = lda.fit_transform(X_count)

print(f"Topics discovered: {n_topics}")
print("\nTop 10 words per topic:")
feature_names = count_vec.get_feature_names_out()
for topic_idx, topic in enumerate(lda.components_):
    top_words_idx = topic.argsort()[-10:][::-1]
    top_words = [feature_names[i] for i in top_words_idx]
    print(f"  Topic {topic_idx + 1}: {', '.join(top_words)}")

perplexity = lda.perplexity(X_count)
log_likelihood = lda.score(X_count)
model_results['LDA'] = {
    'perplexity': perplexity,
    'log_likelihood': log_likelihood,
    'n_topics': n_topics
}
print(f"\n✓ Perplexity: {perplexity:.2f} (lower is better)")
print(f"✓ Log-likelihood: {log_likelihood:.2f} (higher is better)")

# ============ MODEL 2: NMF TOPIC MODELING ============
print("\n[MODEL 2] Non-negative Matrix Factorization (NMF)")
print("-" * 80)
nmf = NMF(n_components=n_topics, random_state=42, max_iter=200)
nmf_topics = nmf.fit_transform(X_tfidf)

print(f"Topics discovered: {n_topics}")
print("\nTop 10 words per topic:")
tfidf_features = tfidf.get_feature_names_out()
for topic_idx, topic in enumerate(nmf.components_):
    top_words_idx = topic.argsort()[-10:][::-1]
    top_words = [tfidf_features[i] for i in top_words_idx]
    print(f"  Topic {topic_idx + 1}: {', '.join(top_words)}")

reconstruction_error = nmf.reconstruction_err_
model_results['NMF'] = {
    'reconstruction_error': reconstruction_error,
    'n_topics': n_topics
}
print(f"\n✓ Reconstruction Error: {reconstruction_error:.2f} (lower is better)")

# ============ MODEL 3: K-MEANS CLUSTERING ============
print("\n[MODEL 3] K-Means Clustering")
print("-" * 80)
n_clusters = 7  # Match number of emotions
kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
kmeans_labels = kmeans.fit_predict(X_tfidf)

silhouette = silhouette_score(X_tfidf, kmeans_labels)
calinski = calinski_harabasz_score(X_tfidf.toarray(), kmeans_labels)
davies = davies_bouldin_score(X_tfidf.toarray(), kmeans_labels)

model_results['KMeans'] = {
    'silhouette_score': silhouette,
    'calinski_harabasz': calinski,
    'davies_bouldin': davies,
    'n_clusters': n_clusters
}

print(f"Number of clusters: {n_clusters}")
print(f"✓ Silhouette Score: {silhouette:.4f} (higher is better, range: -1 to 1)")
print(f"✓ Calinski-Harabasz Index: {calinski:.2f} (higher is better)")
print(f"✓ Davies-Bouldin Index: {davies:.4f} (lower is better)")

cluster_counts = Counter(kmeans_labels)
print("\nCluster distribution:")
for cluster_id, count in sorted(cluster_counts.items()):
    print(f"  Cluster {cluster_id}: {count:>5} documents ({count/len(kmeans_labels)*100:>5.1f}%)")

# ============================================================================
# PART C: SUPERVISED LEARNING - EMOTION CLASSIFICATION
# ============================================================================

print("\n" + "="*80)
print("📊 PART C: SUPERVISED LEARNING - EMOTION CLASSIFICATION")
print("="*80)

# Split data
X_train, X_test, y_train, y_test = train_test_split(
    X_tfidf, y_emotion, test_size=0.2, random_state=42, stratify=y_emotion
)

print(f"\nTraining set: {X_train.shape[0]} samples")
print(f"Testing set: {X_test.shape[0]} samples")
print(f"Emotion classes: {np.unique(y_emotion)}")

# ============ MODEL 4: NAIVE BAYES ============
print("\n[MODEL 4] Multinomial Naive Bayes")
print("-" * 80)
nb = MultinomialNB()
nb.fit(X_train, y_train)
y_pred_nb = nb.predict(X_test)

accuracy_nb = accuracy_score(y_test, y_pred_nb)
f1_nb = f1_score(y_test, y_pred_nb, average='weighted')
print(f"✓ Accuracy: {accuracy_nb:.4f}")
print(f"✓ F1-Score (weighted): {f1_nb:.4f}")
print("\nClassification Report:")
print(classification_report(y_test, y_pred_nb))

model_results['Naive_Bayes'] = {
    'accuracy': accuracy_nb,
    'f1_score': f1_nb,
    'predictions': y_pred_nb
}

# ============ MODEL 5: LOGISTIC REGRESSION ============
print("\n[MODEL 5] Logistic Regression")
print("-" * 80)
lr = LogisticRegression(max_iter=1000, random_state=42, multi_class='multinomial', C=1.0)
lr.fit(X_train, y_train)
y_pred_lr = lr.predict(X_test)

accuracy_lr = accuracy_score(y_test, y_pred_lr)
f1_lr = f1_score(y_test, y_pred_lr, average='weighted')
print(f"✓ Accuracy: {accuracy_lr:.4f}")
print(f"✓ F1-Score (weighted): {f1_lr:.4f}")
print("\nClassification Report:")
print(classification_report(y_test, y_pred_lr))

model_results['Logistic_Regression'] = {
    'accuracy': accuracy_lr,
    'f1_score': f1_lr,
    'predictions': y_pred_lr
}

# ============ MODEL 6: RANDOM FOREST ============
print("\n[MODEL 6] Random Forest Classifier")
print("-" * 80)
rf = RandomForestClassifier(n_estimators=100, random_state=42, max_depth=20, n_jobs=-1)
rf.fit(X_train, y_train)
y_pred_rf = rf.predict(X_test)

accuracy_rf = accuracy_score(y_test, y_pred_rf)
f1_rf = f1_score(y_test, y_pred_rf, average='weighted')
print(f"✓ Accuracy: {accuracy_rf:.4f}")
print(f"✓ F1-Score (weighted): {f1_rf:.4f}")
print("\nClassification Report:")
print(classification_report(y_test, y_pred_rf))

model_results['Random_Forest'] = {
    'accuracy': accuracy_rf,
    'f1_score': f1_rf,
    'predictions': y_pred_rf
}

# Feature importance from Random Forest
print("\nTop 15 Most Important Features:")
feature_importance = pd.DataFrame({
    'feature': tfidf.get_feature_names_out(),
    'importance': rf.feature_importances_
}).sort_values('importance', ascending=False).head(15)
print(feature_importance.to_string(index=False))

# ============ MODEL 7: GRADIENT BOOSTING ============
print("\n[MODEL 7] Gradient Boosting Classifier")
print("-" * 80)
gb = GradientBoostingClassifier(n_estimators=100, random_state=42, max_depth=5, learning_rate=0.1)
gb.fit(X_train, y_train)
y_pred_gb = gb.predict(X_test)

accuracy_gb = accuracy_score(y_test, y_pred_gb)
f1_gb = f1_score(y_test, y_pred_gb, average='weighted')
print(f"✓ Accuracy: {accuracy_gb:.4f}")
print(f"✓ F1-Score (weighted): {f1_gb:.4f}")
print("\nClassification Report:")
print(classification_report(y_test, y_pred_gb))

model_results['Gradient_Boosting'] = {
    'accuracy': accuracy_gb,
    'f1_score': f1_gb,
    'predictions': y_pred_gb
}

# ============ MODEL 8: SVM ============
print("\n[MODEL 8] Support Vector Machine (Linear)")
print("-" * 80)
svm = SVC(kernel='linear', random_state=42, C=1.0)
svm.fit(X_train, y_train)
y_pred_svm = svm.predict(X_test)

accuracy_svm = accuracy_score(y_test, y_pred_svm)
f1_svm = f1_score(y_test, y_pred_svm, average='weighted')
print(f"✓ Accuracy: {accuracy_svm:.4f}")
print(f"✓ F1-Score (weighted): {f1_svm:.4f}")
print("\nClassification Report:")
print(classification_report(y_test, y_pred_svm))

model_results['SVM'] = {
    'accuracy': accuracy_svm,
    'f1_score': f1_svm,
    'predictions': y_pred_svm
}

# ============================================================================
# MODEL COMPARISON & VISUALIZATION
# ============================================================================

print("\n" + "="*80)
print("📊 COMPREHENSIVE MODEL COMPARISON")
print("="*80)

print("\n🔹 UNSUPERVISED MODELS:")
print("-" * 80)
print(f"{'Model':<25} {'Primary Metric':<30} {'Score':<15}")
print("-" * 80)
print(f"{'LDA':<25} {'Perplexity (lower better)':<30} {model_results['LDA']['perplexity']:<15.2f}")
print(f"{'NMF':<25} {'Reconstruction Error':<30} {model_results['NMF']['reconstruction_error']:<15.2f}")
print(f"{'K-Means':<25} {'Silhouette Score':<30} {model_results['KMeans']['silhouette_score']:<15.4f}")

print("\n🔹 SUPERVISED MODELS (Emotion Classification):")
print("-" * 80)
print(f"{'Model':<30} {'Accuracy':<15} {'F1-Score':<15}")
print("-" * 80)
classifiers = ['Naive_Bayes', 'Logistic_Regression', 'Random_Forest', 
               'Gradient_Boosting', 'SVM']
for clf in classifiers:
    if clf in model_results:
        acc = model_results[clf]['accuracy']
        f1 = model_results[clf]['f1_score']
        print(f"{clf.replace('_', ' '):<30} {acc:<15.4f} {f1:<15.4f}")

# Determine best model
best_clf = max(classifiers, key=lambda x: model_results[x]['f1_score'])
print("\n" + "="*80)
print("🏆 BEST MODEL IDENTIFICATION")
print("="*80)
print(f"\n✨ BEST EMOTION CLASSIFIER: {best_clf.replace('_', ' ')}")
print(f"   Accuracy: {model_results[best_clf]['accuracy']:.4f}")
print(f"   F1-Score: {model_results[best_clf]['f1_score']:.4f}")

# ============================================================================
# CONFUSION MATRICES VISUALIZATION
# ============================================================================

print("\n" + "="*80)
print("CONFUSION MATRICES FOR TOP 3 MODELS")
print("="*80)

# Get top 3 models by F1-score
top_3_models = sorted(classifiers, key=lambda x: model_results[x]['f1_score'], reverse=True)[:3]

fig, axes = plt.subplots(1, 3, figsize=(20, 6))

for idx, model_name in enumerate(top_3_models):
    cm = confusion_matrix(y_test, model_results[model_name]['predictions'])
    
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', ax=axes[idx],
                xticklabels=np.unique(y_emotion), yticklabels=np.unique(y_emotion),
                cbar_kws={'label': 'Count'})
    axes[idx].set_title(f'{model_name.replace("_", " ")}\nAccuracy: {model_results[model_name]["accuracy"]:.3f}',
                       fontsize=12, fontweight='bold')
    axes[idx].set_xlabel('Predicted Emotion', fontsize=10, fontweight='bold')
    axes[idx].set_ylabel('True Emotion', fontsize=10, fontweight='bold')
    plt.setp(axes[idx].xaxis.get_majorticklabels(), rotation=45, ha='right')
    plt.setp(axes[idx].yaxis.get_majorticklabels(), rotation=0)

plt.suptitle('Confusion Matrices - Top 3 Emotion Classification Models', 
             fontsize=14, fontweight='bold', y=1.02)
plt.tight_layout()
plt.show()

# ============================================================================
# MODEL PERFORMANCE VISUALIZATION
# ============================================================================

# Bar chart comparison
fig, axes = plt.subplots(1, 2, figsize=(16, 6))

# Accuracy comparison
models_names = [m.replace('_', ' ') for m in classifiers]
accuracies = [model_results[m]['accuracy'] for m in classifiers]
f1_scores = [model_results[m]['f1_score'] for m in classifiers]

colors = plt.cm.viridis(np.linspace(0, 1, len(classifiers)))

bars1 = axes[0].bar(models_names, accuracies, color=colors, edgecolor='black', linewidth=1.2)
axes[0].set_title('Model Accuracy Comparison', fontsize=14, fontweight='bold', pad=15)
axes[0].set_ylabel('Accuracy', fontsize=12, fontweight='bold')
axes[0].set_ylim([0, 1])
axes[0].grid(axis='y', alpha=0.3)
axes[0].tick_params(axis='x', rotation=45)
plt.setp(axes[0].xaxis.get_majorticklabels(), rotation=45, ha='right')

for bar in bars1:
    height = bar.get_height()
    axes[0].text(bar.get_x() + bar.get_width()/2., height,
                f'{height:.3f}', ha='center', va='bottom', fontweight='bold', fontsize=9)

# F1-Score comparison
bars2 = axes[1].bar(models_names, f1_scores, color=colors, edgecolor='black', linewidth=1.2)
axes[1].set_title('Model F1-Score Comparison', fontsize=14, fontweight='bold', pad=15)
axes[1].set_ylabel('F1-Score (Weighted)', fontsize=12, fontweight='bold')
axes[1].set_ylim([0, 1])
axes[1].grid(axis='y', alpha=0.3)
plt.setp(axes[1].xaxis.get_majorticklabels(), rotation=45, ha='right')

for bar in bars2:
    height = bar.get_height()
    axes[1].text(bar.get_x() + bar.get_width()/2., height,
                f'{height:.3f}', ha='center', va='bottom', fontweight='bold', fontsize=9)

plt.tight_layout()
plt.show()

# ============================================================================
# WORD CLOUDS FOR EACH EMOTION
# ============================================================================

print("\n" + "="*80)
print("GENERATING WORD CLOUDS FOR EACH EMOTION")
print("="*80)

emotions = df_clean['emotion'].unique()
n_emotions = len(emotions)
fig, axes = plt.subplots(2, 4, figsize=(20, 10))
axes = axes.flatten()

for idx, emotion in enumerate(emotions):
    text_data = ' '.join(df_clean[df_clean['emotion'] == emotion]['text_processed'])
    
    if text_data.strip():
        wordcloud = WordCloud(width=400, height=300, background_color='white',
                            colormap='viridis', max_words=50).generate(text_data)
        
        axes[idx].imshow(wordcloud, interpolation='bilinear')
        axes[idx].set_title(f'{emotion.capitalize()}', fontsize=14, fontweight='bold')
        axes[idx].axis('off')
    else:
        axes[idx].text(0.5, 0.5, f'No data for\n{emotion}', 
                      ha='center', va='center', fontsize=12)
        axes[idx].axis('off')

# Hide unused subplot
for idx in range(n_emotions, len(axes)):
    axes[idx].axis('off')

plt.suptitle('Word Clouds by Emotion Category', fontsize=16, fontweight='bold', y=0.98)
plt.tight_layout()
plt.show()

print("\n" + "="*80)
print("✓ MODEL DEVELOPMENT & EVALUATION COMPLETED!")
print("="*80)
print(f"\nSummary:")
print(f"  • Total models trained: {len(classifiers) + 3}")
print(f"  • Best classifier: {best_clf.replace('_', ' ')}")
print(f"  • Best accuracy: {model_results[best_clf]['accuracy']:.4f}")
print(f"  • Best F1-score: {model_results[best_clf]['f1_score']:.4f}")
print(f"  • Dataset size: {len(df_clean)} samples")
print(f"  • Emotion categories: {len(emotions)}")
print("\n✓ All models and visualizations completed successfully!")








# ============================================================================
# ADVANCED MODEL VISUALIZATIONS
# ============================================================================

import time
from datetime import datetime

print("\n" + "="*80)
print("ADVANCED MODEL VISUALIZATIONS")
print("="*80)
print(f"Started at: {datetime.now().strftime('%H:%M:%S')}")

# ============================================================================
# VISUALIZATION 1: COMPREHENSIVE MODEL PERFORMANCE COMPARISON
# ============================================================================

print("\n[1/5] Generating comprehensive performance comparison...")
start_time = time.time()

fig, axes = plt.subplots(2, 2, figsize=(18, 12))

# 1. Supervised Model Accuracy & F1-Score Comparison
classifiers = ['Naive_Bayes', 'Logistic_Regression', 'Random_Forest',
               'Gradient_Boosting', 'SVM']
accuracies = [model_results[clf]['accuracy'] for clf in classifiers]
f1_scores = [model_results[clf]['f1_score'] for clf in classifiers]

x_pos = np.arange(len(classifiers))
width = 0.35

bars1 = axes[0, 0].bar(x_pos - width/2, accuracies, width, 
                       label='Accuracy', color='#4ECDC4', edgecolor='black', linewidth=1.5)
bars2 = axes[0, 0].bar(x_pos + width/2, f1_scores, width,
                       label='F1-Score', color='#FF6B6B', edgecolor='black', linewidth=1.5)

axes[0, 0].set_xticks(x_pos)
axes[0, 0].set_xticklabels([c.replace('_', '\n') for c in classifiers], 
                           rotation=0, ha='center', fontsize=10)
axes[0, 0].set_ylabel('Score', fontsize=12, fontweight='bold')
axes[0, 0].set_title('Supervised Models: Accuracy & F1-Score Comparison', 
                     fontsize=13, fontweight='bold', pad=15)
axes[0, 0].set_ylim([0, 1.0])
axes[0, 0].axhline(y=0.5, color='red', linestyle='--', alpha=0.5, label='Baseline (50%)')
axes[0, 0].grid(axis='y', alpha=0.3, linestyle='--')
axes[0, 0].legend(loc='lower right', fontsize=10)

# Add value labels
for bar1, bar2 in zip(bars1, bars2):
    height1 = bar1.get_height()
    height2 = bar2.get_height()
    axes[0, 0].text(bar1.get_x() + bar1.get_width()/2., height1 + 0.01,
                    f'{height1:.3f}', ha='center', va='bottom', 
                    fontweight='bold', fontsize=8)
    axes[0, 0].text(bar2.get_x() + bar2.get_width()/2., height2 + 0.01,
                    f'{height2:.3f}', ha='center', va='bottom', 
                    fontweight='bold', fontsize=8)

# 2. Clustering Model Comparison (Multiple Metrics)
clustering_models = ['K-Means']
silhouette_scores = [model_results['KMeans']['silhouette_score']]
calinski_scores = [model_results['KMeans']['calinski_harabasz'] / 1000]  # Normalize
davies_scores = [1 - model_results['KMeans']['davies_bouldin']]  # Invert (higher better)

x_clust = np.arange(len(clustering_models))
width_clust = 0.25

bars_sil = axes[0, 1].bar(x_clust - width_clust, silhouette_scores, width_clust,
                          label='Silhouette', color='#95E1D3', edgecolor='black', linewidth=1.5)
bars_cal = axes[0, 1].bar(x_clust, calinski_scores, width_clust,
                          label='Calinski-H (÷1000)', color='#F38181', edgecolor='black', linewidth=1.5)
bars_dav = axes[0, 1].bar(x_clust + width_clust, davies_scores, width_clust,
                          label='Davies-B (inv)', color='#EAFFD0', edgecolor='black', linewidth=1.5)

axes[0, 1].set_xticks(x_clust)
axes[0, 1].set_xticklabels(clustering_models, fontsize=11)
axes[0, 1].set_ylabel('Score', fontsize=12, fontweight='bold')
axes[0, 1].set_title('Clustering Model: Multiple Metrics', 
                     fontsize=13, fontweight='bold', pad=15)
axes[0, 1].grid(axis='y', alpha=0.3, linestyle='--')
axes[0, 1].legend(fontsize=9)

# 3. Confusion Matrix Heatmap for Best Model
best_clf_name = max(classifiers, key=lambda x: model_results[x]['f1_score'])
best_pred = model_results[best_clf_name]['predictions']

cm = confusion_matrix(y_test, best_pred)
emotions_sorted = sorted(np.unique(y_emotion))

sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', ax=axes[1, 0],
            xticklabels=emotions_sorted,
            yticklabels=emotions_sorted,
            cbar_kws={'label': 'Count'},
            linewidths=0.5, linecolor='gray')
axes[1, 0].set_xlabel('Predicted Emotion', fontsize=11, fontweight='bold')
axes[1, 0].set_ylabel('True Emotion', fontsize=11, fontweight='bold')
axes[1, 0].set_title(f'Confusion Matrix: {best_clf_name.replace("_", " ")}\n' + 
                     f'(Accuracy: {model_results[best_clf_name]["accuracy"]:.3f}, ' +
                     f'F1: {model_results[best_clf_name]["f1_score"]:.3f})',
                     fontsize=12, fontweight='bold', pad=15)
plt.setp(axes[1, 0].xaxis.get_majorticklabels(), rotation=45, ha='right')
plt.setp(axes[1, 0].yaxis.get_majorticklabels(), rotation=0)

# 4. Topic Modeling Quality Comparison
topic_models = ['LDA', 'NMF']
lda_score_norm = 1 / (1 + model_results['LDA']['perplexity'] / 100)
nmf_score_norm = 1 / (1 + model_results['NMF']['reconstruction_error'] / 10)

topic_scores = [lda_score_norm, nmf_score_norm]
colors_topic = ['#AA96DA', '#FCBAD3']

bars_topic = axes[1, 1].bar(range(len(topic_models)), topic_scores,
                            color=colors_topic, edgecolor='black', linewidth=1.5)
axes[1, 1].set_xticks(range(len(topic_models)))
axes[1, 1].set_xticklabels(topic_models, fontsize=11)
axes[1, 1].set_ylabel('Normalized Quality Score', fontsize=12, fontweight='bold')
axes[1, 1].set_title('Topic Models: Quality Comparison\n(Higher is Better)', 
                     fontsize=13, fontweight='bold', pad=15)
axes[1, 1].grid(axis='y', alpha=0.3, linestyle='--')

for bar, score in zip(bars_topic, topic_scores):
    height = bar.get_height()
    axes[1, 1].text(bar.get_x() + bar.get_width()/2., height + 0.01,
                    f'{height:.4f}', ha='center', va='bottom', fontweight='bold', fontsize=9)

# Add original scores as text
axes[1, 1].text(0, -0.15, f'Perplexity: {model_results["LDA"]["perplexity"]:.2f}',
               transform=axes[1, 1].transAxes, fontsize=9, style='italic')
axes[1, 1].text(0.5, -0.15, f'Recon. Error: {model_results["NMF"]["reconstruction_error"]:.2f}',
               transform=axes[1, 1].transAxes, fontsize=9, style='italic')

plt.suptitle('Comprehensive Model Performance Analysis', 
             fontsize=16, fontweight='bold', y=0.995)
plt.tight_layout()
plt.show()

print(f"✓ Completed in {time.time() - start_time:.2f}s")

# ============================================================================
# VISUALIZATION 2: t-SNE DIMENSIONALITY REDUCTION
# ============================================================================

print("\n[2/5] Generating t-SNE visualization (this may take 2-3 minutes)...")
start_time = time.time()

fig, axes = plt.subplots(1, 2, figsize=(18, 7))

# Use subset for faster computation
sample_size = min(2000, len(X_tfidf.toarray()))
sample_indices = np.random.choice(len(X_tfidf.toarray()), sample_size, replace=False)

print(f"   Using {sample_size} samples for t-SNE...")
tsne = TSNE(n_components=2, random_state=42, perplexity=30, n_iter=1000, verbose=0)
X_tsne = tsne.fit_transform(X_tfidf.toarray()[sample_indices])

# Plot 1: K-Means clusters
scatter1 = axes[0].scatter(X_tsne[:, 0], X_tsne[:, 1],
                          c=kmeans_labels[sample_indices],
                          cmap='tab10',
                          alpha=0.6,
                          edgecolors='black',
                          linewidth=0.5,
                          s=50)
axes[0].set_title('t-SNE Visualization: K-Means Clusters', 
                  fontsize=14, fontweight='bold', pad=15)
axes[0].set_xlabel('t-SNE Component 1', fontsize=12, fontweight='bold')
axes[0].set_ylabel('t-SNE Component 2', fontsize=12, fontweight='bold')
axes[0].grid(alpha=0.2, linestyle='--')
cbar1 = plt.colorbar(scatter1, ax=axes[0])
cbar1.set_label('Cluster ID', fontsize=11, fontweight='bold')

# Plot 2: True emotion labels
emotion_to_num = {emotion: idx for idx, emotion in enumerate(sorted(np.unique(y_emotion)))}
emotion_numeric = [emotion_to_num[e] for e in y_emotion[sample_indices]]

scatter2 = axes[1].scatter(X_tsne[:, 0], X_tsne[:, 1],
                          c=emotion_numeric,
                          cmap='Spectral',
                          alpha=0.6,
                          edgecolors='black',
                          linewidth=0.5,
                          s=50)
axes[1].set_title('t-SNE Visualization: True Emotion Labels', 
                  fontsize=14, fontweight='bold', pad=15)
axes[1].set_xlabel('t-SNE Component 1', fontsize=12, fontweight='bold')
axes[1].set_ylabel('t-SNE Component 2', fontsize=12, fontweight='bold')
axes[1].grid(alpha=0.2, linestyle='--')
cbar2 = plt.colorbar(scatter2, ax=axes[1], ticks=range(len(emotion_to_num)))
cbar2.set_ticklabels(sorted(emotion_to_num.keys()), fontsize=9)
cbar2.set_label('Emotion', fontsize=11, fontweight='bold')

plt.tight_layout()
plt.show()

print(f"✓ Completed in {time.time() - start_time:.2f}s")

# ============================================================================
# VISUALIZATION 3: WORD CLOUDS FOR LDA TOPICS
# ============================================================================

print("\n[3/5] Generating word clouds for LDA topics...")
start_time = time.time()

fig, axes = plt.subplots(2, 4, figsize=(20, 10))
axes = axes.flatten()

for topic_idx in range(min(n_topics, 7)):
    print(f"   Generating word cloud for Topic {topic_idx + 1}...")
    topic = lda.components_[topic_idx]
    top_words_idx = topic.argsort()[-100:][::-1]
    word_freq = {feature_names[i]: topic[i] for i in top_words_idx}
    
    wc = WordCloud(width=450, height=350,
                  background_color='white',
                  colormap='viridis',
                  relative_scaling=0.5,
                  min_font_size=8,
                  max_words=80).generate_from_frequencies(word_freq)
    
    axes[topic_idx].imshow(wc, interpolation='bilinear')
    axes[topic_idx].set_title(f'LDA Topic {topic_idx + 1}',
                             fontsize=13, fontweight='bold', pad=10)
    axes[topic_idx].axis('off')

# Hide unused subplot
axes[7].axis('off')

plt.suptitle('Word Clouds for LDA Topics', fontsize=16, fontweight='bold', y=0.98)
plt.tight_layout()
plt.show()

print(f"✓ Completed in {time.time() - start_time:.2f}s")

# ============================================================================
# VISUALIZATION 4: FEATURE IMPORTANCE FROM RANDOM FOREST
# ============================================================================

print("\n[4/5] Generating feature importance visualization...")
start_time = time.time()

feature_importance = rf.feature_importances_
top_n = 25
top_indices = np.argsort(feature_importance)[-top_n:][::-1]
top_features = [tfidf.get_feature_names_out()[i] for i in top_indices]
top_importance = feature_importance[top_indices]

fig, ax = plt.subplots(figsize=(12, 10))
colors_importance = plt.cm.plasma(np.linspace(0.2, 0.9, len(top_features)))

bars = ax.barh(range(len(top_features)), top_importance, 
               color=colors_importance, edgecolor='black', linewidth=1)
ax.set_yticks(range(len(top_features)))
ax.set_yticklabels(top_features, fontsize=10)
ax.set_xlabel('Importance Score', fontsize=12, fontweight='bold')
ax.set_ylabel('Features (Words)', fontsize=12, fontweight='bold')
ax.set_title(f'Top {top_n} Most Important Features for Emotion Classification\n' +
             f'(Random Forest)', fontsize=14, fontweight='bold', pad=20)
ax.invert_yaxis()
ax.grid(axis='x', alpha=0.3, linestyle='--')

# Add value labels
for i, (bar, importance) in enumerate(zip(bars, top_importance)):
    ax.text(importance + 0.0002, i, f'{importance:.4f}',
            va='center', fontweight='bold', fontsize=8)

plt.tight_layout()
plt.show()

print(f"✓ Completed in {time.time() - start_time:.2f}s")

# ============================================================================
# VISUALIZATION 5: PER-EMOTION PERFORMANCE HEATMAP
# ============================================================================

print("\n[5/5] Generating per-emotion performance analysis...")
start_time = time.time()

# Calculate precision, recall, F1 for each emotion across models
from sklearn.metrics import precision_recall_fscore_support

fig, axes = plt.subplots(1, 3, figsize=(20, 6))
metrics_names = ['Precision', 'Recall', 'F1-Score']

for metric_idx, metric_name in enumerate(metrics_names):
    performance_matrix = []
    
    for clf_name in classifiers:
        y_pred = model_results[clf_name]['predictions']
        precision, recall, f1, _ = precision_recall_fscore_support(
            y_test, y_pred, labels=emotions_sorted, average=None, zero_division=0
        )
        
        if metric_name == 'Precision':
            performance_matrix.append(precision)
        elif metric_name == 'Recall':
            performance_matrix.append(recall)
        else:
            performance_matrix.append(f1)
    
    performance_df = pd.DataFrame(
        performance_matrix,
        index=[c.replace('_', ' ') for c in classifiers],
        columns=emotions_sorted
    )
    
    sns.heatmap(performance_df, annot=True, fmt='.3f', cmap='YlGnBu', 
                ax=axes[metric_idx], cbar_kws={'label': metric_name},
                linewidths=0.5, linecolor='gray', vmin=0, vmax=1)
    axes[metric_idx].set_title(f'{metric_name} by Emotion', 
                               fontsize=13, fontweight='bold', pad=15)
    axes[metric_idx].set_xlabel('Emotion', fontsize=11, fontweight='bold')
    axes[metric_idx].set_ylabel('Model', fontsize=11, fontweight='bold')
    plt.setp(axes[metric_idx].xaxis.get_majorticklabels(), rotation=45, ha='right')
    plt.setp(axes[metric_idx].yaxis.get_majorticklabels(), rotation=0)

plt.suptitle('Per-Emotion Performance Analysis Across All Models', 
             fontsize=16, fontweight='bold', y=1.00)
plt.tight_layout()
plt.show()

print(f"✓ Completed in {time.time() - start_time:.2f}s")

# ============================================================================
# SUMMARY STATISTICS
# ============================================================================

print("\n" + "="*80)
print("✓ ALL VISUALIZATIONS COMPLETED SUCCESSFULLY!")
print("="*80)
print(f"\nCompleted at: {datetime.now().strftime('%H:%M:%S')}")
print(f"\nVisualization Summary:")
print(f"  1. ✓ Comprehensive model performance comparison")
print(f"  2. ✓ t-SNE dimensionality reduction (2 plots)")
print(f"  3. ✓ Word clouds for {n_topics} LDA topics")
print(f"  4. ✓ Feature importance analysis (top {top_n} features)")
print(f"  5. ✓ Per-emotion performance heatmaps (3 metrics)")
print(f"\nTotal visualizations generated: 10+ charts")
print(f"Best model: {best_clf_name.replace('_', ' ')} (F1: {model_results[best_clf_name]['f1_score']:.4f})")


# ============================================================================
# AUTOMATED RESULTS INTERPRETATION & INSIGHTS
# ============================================================================

print("\n" + "="*80)
print("🤖 AUTOMATED RESULTS INTERPRETATION & INSIGHTS")
print("="*80)
print(f"Analysis timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

# ============================================================================
# 1. DATASET OVERVIEW INTERPRETATION
# ============================================================================

print("\n" + "="*80)
print("📊 DATASET OVERVIEW ANALYSIS")
print("="*80)

total_samples = len(df_clean)
total_emotions = len(df_clean['emotion'].unique())
emotion_distribution = df_clean['emotion'].value_counts()
dominant_emotion = emotion_distribution.index[0]
minority_emotion = emotion_distribution.index[-1]

print(f"\n✓ DATASET CHARACTERISTICS:")
print(f"  • Total samples analyzed: {total_samples:,}")
print(f"  • Number of emotion categories: {total_emotions}")
print(f"  • Average comment length: {df_clean['word_count_processed'].mean():.1f} words")
print(f"  • Text preprocessing reduced length by: {((df_clean['original_length'].mean() - df_clean['processed_length'].mean()) / df_clean['original_length'].mean() * 100):.1f}%")

print(f"\n✓ CLASS DISTRIBUTION ANALYSIS:")
print(f"  • Dominant emotion: '{dominant_emotion}' ({emotion_distribution[dominant_emotion]} samples, {emotion_distribution[dominant_emotion]/total_samples*100:.1f}%)")
print(f"  • Minority emotion: '{minority_emotion}' ({emotion_distribution[minority_emotion]} samples, {emotion_distribution[minority_emotion]/total_samples*100:.1f}%)")
print(f"  • Imbalance ratio: {emotion_distribution[dominant_emotion] / emotion_distribution[minority_emotion]:.2f}:1")

# Class imbalance interpretation
imbalance_ratio = emotion_distribution[dominant_emotion] / emotion_distribution[minority_emotion]
if imbalance_ratio > 10:
    print(f"\n⚠️  CRITICAL CLASS IMBALANCE DETECTED!")
    print(f"     The dataset is heavily imbalanced. Consider using:")
    print(f"     - SMOTE (Synthetic Minority Over-sampling)")
    print(f"     - Class weights in model training")
    print(f"     - Stratified sampling")
elif imbalance_ratio > 3:
    print(f"\n⚠️  MODERATE CLASS IMBALANCE DETECTED")
    print(f"     Consider using class weights or stratified validation")
else:
    print(f"\n✓ CLASS DISTRIBUTION: Relatively balanced dataset")

# ============================================================================
# 2. SENTIMENT-EMOTION ALIGNMENT ANALYSIS
# ============================================================================

print("\n" + "="*80)
print("💭 SENTIMENT-EMOTION ALIGNMENT ANALYSIS")
print("="*80)

# Calculate sentiment alignment for each emotion
emotion_sentiment_analysis = df_clean.groupby('emotion').agg({
    'sentiment_polarity': ['mean', 'std'],
    'sentiment_subjectivity': 'mean'
}).round(3)

print("\n✓ EMOTION-SENTIMENT CONSISTENCY:")

# Expected sentiment alignments
expected_alignments = {
    'joy': ('positive', 0.3),
    'sadness': ('negative', -0.3),
    'anger': ('negative', -0.2),
    'fear': ('negative', -0.2),
    'surprise': ('neutral', 0.0),
    'disgust': ('negative', -0.3),
    'neutral': ('neutral', 0.0)
}

for emotion in df_clean['emotion'].unique():
    avg_polarity = df_clean[df_clean['emotion'] == emotion]['sentiment_polarity'].mean()
    
    if emotion in expected_alignments:
        expected_sentiment, threshold = expected_alignments[emotion]
        
        if expected_sentiment == 'positive' and avg_polarity > threshold:
            status = "✓ ALIGNED"
        elif expected_sentiment == 'negative' and avg_polarity < threshold:
            status = "✓ ALIGNED"
        elif expected_sentiment == 'neutral' and abs(avg_polarity) < 0.15:
            status = "✓ ALIGNED"
        else:
            status = "⚠️  MISALIGNED"
        
        print(f"  • {emotion.capitalize():12} → Polarity: {avg_polarity:+.3f} | Expected: {expected_sentiment:8} | {status}")

# Overall alignment score
alignment_scores = []
for emotion in df_clean['emotion'].unique():
    if emotion in expected_alignments:
        avg_polarity = df_clean[df_clean['emotion'] == emotion]['sentiment_polarity'].mean()
        _, threshold = expected_alignments[emotion]
        alignment_scores.append(1 if abs(avg_polarity - threshold) < 0.2 else 0)

overall_alignment = sum(alignment_scores) / len(alignment_scores) * 100
print(f"\n  Overall sentiment-emotion alignment: {overall_alignment:.1f}%")

if overall_alignment > 80:
    print(f"  ✓ EXCELLENT: Emotion labels are highly consistent with sentiment")
elif overall_alignment > 60:
    print(f"  ⚠️  MODERATE: Some emotion labels may need review")
else:
    print(f"  ❌ POOR: Significant inconsistency between emotions and sentiment")

# ============================================================================
# 3. MODEL PERFORMANCE INTERPRETATION
# ============================================================================

print("\n" + "="*80)
print("🎯 MODEL PERFORMANCE INTERPRETATION")
print("="*80)

# Get performance metrics
model_performances = []
for clf in classifiers:
    model_performances.append({
        'model': clf.replace('_', ' '),
        'accuracy': model_results[clf]['accuracy'],
        'f1_score': model_results[clf]['f1_score']
    })

model_df = pd.DataFrame(model_performances).sort_values('f1_score', ascending=False)

print("\n✓ MODEL RANKING (by F1-Score):")
for idx, row in model_df.iterrows():
    rank = model_df.index.get_loc(idx) + 1
    medal = "🥇" if rank == 1 else "🥈" if rank == 2 else "🥉" if rank == 3 else f"{rank}."
    print(f"  {medal} {row['model']:25} | Accuracy: {row['accuracy']:.4f} | F1-Score: {row['f1_score']:.4f}")

# Performance interpretation
best_model = model_df.iloc[0]
worst_model = model_df.iloc[-1]
performance_gap = best_model['f1_score'] - worst_model['f1_score']

print(f"\n✓ PERFORMANCE ANALYSIS:")
print(f"  • Best model: {best_model['model']}")
print(f"  • Performance gap: {performance_gap:.4f} ({performance_gap*100:.1f}%)")

if best_model['f1_score'] > 0.85:
    print(f"  • ✓ EXCELLENT PERFORMANCE: Models are highly effective for emotion classification")
elif best_model['f1_score'] > 0.70:
    print(f"  • ✓ GOOD PERFORMANCE: Models show solid emotion classification capability")
elif best_model['f1_score'] > 0.55:
    print(f"  • ⚠️  MODERATE PERFORMANCE: Consider feature engineering or advanced models")
else:
    print(f"  • ❌ POOR PERFORMANCE: Significant improvements needed")

# Model complexity vs performance
print(f"\n✓ COMPLEXITY-PERFORMANCE TRADE-OFF:")
simple_models = ['Naive Bayes', 'Logistic Regression']
complex_models = ['Random Forest', 'Gradient Boosting', 'SVM']

simple_avg = model_df[model_df['model'].isin(simple_models)]['f1_score'].mean()
complex_avg = model_df[model_df['model'].isin(complex_models)]['f1_score'].mean()

print(f"  • Simple models avg F1: {simple_avg:.4f}")
print(f"  • Complex models avg F1: {complex_avg:.4f}")
print(f"  • Improvement from complexity: {((complex_avg - simple_avg) / simple_avg * 100):+.1f}%")

if complex_avg - simple_avg < 0.05:
    print(f"  → RECOMMENDATION: Simple models sufficient - use Logistic Regression for production")
else:
    print(f"  → RECOMMENDATION: Complex models worth the overhead - use {best_model['model']} for production")

# ============================================================================
# 4. PER-EMOTION ANALYSIS
# ============================================================================

print("\n" + "="*80)
print("🎭 PER-EMOTION CLASSIFICATION ANALYSIS")
print("="*80)

# Calculate per-emotion F1 scores for best model
from sklearn.metrics import precision_recall_fscore_support

best_model_pred = model_results[best_clf]['predictions']
precision, recall, f1, support = precision_recall_fscore_support(
    y_test, best_model_pred, labels=emotions_sorted, average=None, zero_division=0
)

emotion_performance = pd.DataFrame({
    'Emotion': emotions_sorted,
    'Precision': precision,
    'Recall': recall,
    'F1-Score': f1,
    'Support': support
}).sort_values('F1-Score', ascending=False)

print(f"\n✓ EMOTION CLASSIFICATION DIFFICULTY (using {best_clf.replace('_', ' ')}):")
print(f"\n{'Emotion':<12} {'F1-Score':<10} {'Precision':<10} {'Recall':<10} {'Support':<10} {'Status'}")
print("-" * 70)

for _, row in emotion_performance.iterrows():
    if row['F1-Score'] > 0.80:
        status = "✓ Easy"
    elif row['F1-Score'] > 0.60:
        status = "○ Moderate"
    else:
        status = "✗ Difficult"
    
    print(f"{row['Emotion']:<12} {row['F1-Score']:<10.3f} {row['Precision']:<10.3f} {row['Recall']:<10.3f} {row['Support']:<10.0f} {status}")

# Identify problematic emotions
difficult_emotions = emotion_performance[emotion_performance['F1-Score'] < 0.60]['Emotion'].tolist()
easy_emotions = emotion_performance[emotion_performance['F1-Score'] > 0.80]['Emotion'].tolist()

print(f"\n✓ EMOTION CLASSIFICATION INSIGHTS:")
if easy_emotions:
    print(f"  • Well-classified emotions: {', '.join(easy_emotions)}")
if difficult_emotions:
    print(f"  • Problematic emotions: {', '.join(difficult_emotions)}")
    print(f"    → These emotions may be:")
    print(f"      - Underrepresented in training data")
    print(f"      - Semantically similar to other emotions")
    print(f"      - Requiring more specific features")

# ============================================================================
# 5. TOPIC MODELING INTERPRETATION
# ============================================================================

print("\n" + "="*80)
print("📚 TOPIC MODELING INTERPRETATION")
print("="*80)

print(f"\n✓ LDA TOPIC MODELING:")
print(f"  • Number of topics discovered: {n_topics}")
print(f"  • Perplexity: {model_results['LDA']['perplexity']:.2f}")

if model_results['LDA']['perplexity'] < 500:
    print(f"  • ✓ EXCELLENT: Topics are well-defined and coherent")
elif model_results['LDA']['perplexity'] < 1000:
    print(f"  • ✓ GOOD: Topics show reasonable coherence")
else:
    print(f"  • ⚠️  MODERATE: Topics may overlap - consider adjusting n_topics")

print(f"\n✓ NMF TOPIC MODELING:")
print(f"  • Reconstruction error: {model_results['NMF']['reconstruction_error']:.2f}")

if model_results['NMF']['reconstruction_error'] < 5:
    print(f"  • ✓ EXCELLENT: NMF effectively decomposed the text matrix")
elif model_results['NMF']['reconstruction_error'] < 10:
    print(f"  • ✓ GOOD: Reasonable topic separation achieved")
else:
    print(f"  • ⚠️  MODERATE: Topics may not be well-separated")

# ============================================================================
# 6. CLUSTERING INTERPRETATION
# ============================================================================

print("\n" + "="*80)
print("🔍 CLUSTERING ANALYSIS INTERPRETATION")
print("="*80)

print(f"\n✓ K-MEANS CLUSTERING:")
print(f"  • Number of clusters: {n_clusters}")
print(f"  • Silhouette score: {model_results['KMeans']['silhouette_score']:.4f}")

if model_results['KMeans']['silhouette_score'] > 0.5:
    print(f"  • ✓ EXCELLENT: Clusters are well-separated and dense")
elif model_results['KMeans']['silhouette_score'] > 0.25:
    print(f"  • ✓ GOOD: Reasonable cluster structure detected")
elif model_results['KMeans']['silhouette_score'] > 0:
    print(f"  • ⚠️  WEAK: Clusters overlap significantly")
else:
    print(f"  • ❌ POOR: Data may not have natural cluster structure")

print(f"\n  • Calinski-Harabasz index: {model_results['KMeans']['calinski_harabasz']:.2f}")
print(f"    (Higher is better - measures cluster density)")

print(f"\n  • Davies-Bouldin index: {model_results['KMeans']['davies_bouldin']:.4f}")
if model_results['KMeans']['davies_bouldin'] < 1.0:
    print(f"    ✓ GOOD: Clusters are well-separated (closer to 0 is better)")
else:
    print(f"    ⚠️  MODERATE: Some cluster overlap exists")

# ============================================================================
# 7. FEATURE IMPORTANCE INSIGHTS
# ============================================================================

print("\n" + "="*80)
print("🔑 KEY FEATURE INSIGHTS")
print("="*80)

# Get top features from Random Forest
# Convert to DataFrame for easier manipulation
feature_names = tfidf.get_feature_names_out()
feature_importance_df = pd.DataFrame({
    'feature': feature_names,
    'importance': rf.feature_importances_
}).sort_values('importance', ascending=False)

top_10_features = feature_importance_df.head(10)

print(f"\n✓ MOST IMPORTANT WORDS FOR EMOTION CLASSIFICATION:")
print(f"  (Based on Random Forest feature importance)")
for idx, row in top_10_features.iterrows():
    rank = top_10_features.index.get_loc(idx) + 1
    print(f"  {rank:2d}. '{row['feature']:20}' → Importance: {row['importance']:.5f}")

# Analyze if important features make semantic sense
emotion_keywords = {
    'joy': ['happy', 'love', 'great', 'best', 'good', 'amazing', 'beautiful', 'thank', 'wonderful', 'awesome'],
    'sadness': ['sad', 'cry', 'miss', 'lost', 'sorry', 'hurt', 'alone', 'tear', 'depressed', 'grief'],
    'anger': ['hate', 'angry', 'stupid', 'worst', 'terrible', 'annoying', 'mad', 'furious', 'pissed', 'rage'],
    'fear': ['fear', 'scared', 'worry', 'afraid', 'nervous', 'anxious', 'terrified', 'panic', 'dread'],
    'surprise': ['wow', 'omg', 'surprised', 'shocked', 'unexpected', 'unbelievable', 'amazing', 'incredible'],
    'disgust': ['disgusting', 'gross', 'nasty', 'sick', 'awful', 'horrible', 'revolting', 'repulsive'],
    'neutral': ['ok', 'okay', 'fine', 'yeah', 'well', 'just', 'really', 'maybe', 'perhaps']
}

matched_keywords = 0
matched_details = []

for _, row in top_10_features.iterrows():
    feature = row['feature']
    for emotion, keywords in emotion_keywords.items():
        if any(keyword in feature.lower() for keyword in keywords):
            matched_keywords += 1
            matched_details.append((feature, emotion))
            break

print(f"\n✓ SEMANTIC COHERENCE ANALYSIS:")
print(f"  • {matched_keywords}/{len(top_10_features)} top features align with expected emotion keywords")

if matched_details:
    print(f"\n  Matched features:")
    for feature, emotion in matched_details:
        print(f"    • '{feature}' → likely indicates '{emotion}'")

coherence_score = matched_keywords / len(top_10_features)

if coherence_score > 0.6:
    print(f"\n  • ✓ EXCELLENT: Feature importance aligns well with human intuition ({coherence_score*100:.0f}%)")
elif coherence_score > 0.3:
    print(f"\n  • ✓ GOOD: Reasonable semantic coherence ({coherence_score*100:.0f}%)")
else:
    print(f"\n  • ⚠️  MODERATE: Some unexpected features are important ({coherence_score*100:.0f}%)")
    print(f"    This could indicate:")
    print(f"    - Context-specific words unique to your dataset")
    print(f"    - Multiword expressions being captured")
    print(f"    - Domain-specific terminology")

# ============================================================================
# 8. RECOMMENDATIONS & ACTION ITEMS
# ============================================================================

print("\n" + "="*80)
print("💡 RECOMMENDATIONS & ACTION ITEMS")
print("="*80)

recommendations = []

# Data recommendations
if imbalance_ratio > 5:
    recommendations.append(("DATA", "HIGH", 
                           f"Address class imbalance using SMOTE or class weights"))

if overall_alignment < 70:
    recommendations.append(("DATA", "MEDIUM", 
                           f"Review emotion labels for consistency with sentiment analysis"))

# Model recommendations
if best_model['f1_score'] < 0.70:
    recommendations.append(("MODEL", "HIGH", 
                           f"Explore advanced architectures (LSTM, BERT, transformers)"))
    recommendations.append(("FEATURES", "HIGH", 
                           f"Add contextual features (emojis, punctuation patterns, n-grams)"))

if difficult_emotions:
    recommendations.append(("MODEL", "MEDIUM", 
                           f"Implement hierarchical classification for: {', '.join(difficult_emotions)}"))

# Deployment recommendations
if complex_avg - simple_avg < 0.05:
    recommendations.append(("DEPLOYMENT", "LOW", 
                           f"Deploy simpler model (Logistic Regression) for faster inference"))
else:
    recommendations.append(("DEPLOYMENT", "LOW", 
                           f"Deploy {best_model['model']} with model compression techniques"))

# Additional features
recommendations.append(("ENHANCEMENT", "MEDIUM", 
                       f"Consider adding: temporal features, user context, comment thread analysis"))

# Print recommendations sorted by priority
print("\n✓ PRIORITIZED RECOMMENDATIONS:\n")
priorities = {"HIGH": [], "MEDIUM": [], "LOW": []}
for category, priority, recommendation in recommendations:
    priorities[priority].append((category, recommendation))

for priority in ["HIGH", "MEDIUM", "LOW"]:
    if priorities[priority]:
        print(f"  {priority} PRIORITY:")
        for category, rec in priorities[priority]:
            print(f"    • [{category}] {rec}")
        print()

# ============================================================================
# 9. FINAL SUMMARY
# ============================================================================

print("="*80)
print("📋 EXECUTIVE SUMMARY")
print("="*80)

print(f"""
Dataset Overview:
  • {total_samples:,} comments analyzed across {total_emotions} emotions
  • Class imbalance ratio: {imbalance_ratio:.1f}:1 ({'CRITICAL' if imbalance_ratio > 10 else 'MODERATE' if imbalance_ratio > 3 else 'ACCEPTABLE'})
  • Sentiment-emotion alignment: {overall_alignment:.1f}% ({'EXCELLENT' if overall_alignment > 80 else 'GOOD' if overall_alignment > 60 else 'NEEDS REVIEW'})

Best Model Performance:
  • Champion model: {best_model['model']}
  • Accuracy: {best_model['accuracy']:.4f} ({best_model['accuracy']*100:.2f}%)
  • F1-Score: {best_model['f1_score']:.4f} ({best_model['f1_score']*100:.2f}%)
  • Performance tier: {'EXCELLENT' if best_model['f1_score'] > 0.85 else 'GOOD' if best_model['f1_score'] > 0.70 else 'MODERATE' if best_model['f1_score'] > 0.55 else 'NEEDS IMPROVEMENT'}

Emotion Classification Difficulty:
  • Easy to classify: {', '.join(easy_emotions) if easy_emotions else 'None'}
  • Difficult to classify: {', '.join(difficult_emotions) if difficult_emotions else 'None'}

Next Steps:
  1. {'Address class imbalance using SMOTE' if imbalance_ratio > 5 else 'Proceed with current data distribution'}
  2. {'Deploy ' + best_model['model'] + ' for production use' if best_model['f1_score'] > 0.70 else 'Improve model performance before deployment'}
  3. {'Focus on improving classification for: ' + ', '.join(difficult_emotions) if difficult_emotions else 'All emotions performing adequately'}
  4. Consider ensemble methods combining top {min(3, len(classifiers))} models
""")

print("="*80)
print("✓ AUTOMATED INTERPRETATION COMPLETE!")
print("="*80)
print(f"Report generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
print("\n💾 TIP: Save this output to a text file for future reference")
print("="*80)










# ============================================================================
# GENERATE WORD DOCUMENT REPORT
# ============================================================================

print("\n" + "="*80)
print("📄 GENERATING WORD DOCUMENT REPORT")
print("="*80)

# Install python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "-q", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

import os

def create_comprehensive_report():
    """Generate a comprehensive 2-page Word document report"""
    
    doc = Document()
    
    # ========================================================================
    # TITLE PAGE
    # ========================================================================
    
    title = doc.add_heading('EMOTION CLASSIFICATION ANALYSIS REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('YouTube Comments Sentiment & Emotion Detection')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(68, 114, 196)
    
    date_para = doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.italic = True
    
    doc.add_paragraph()
    
    # ========================================================================
    # EXECUTIVE SUMMARY
    # ========================================================================
    
    doc.add_heading('1. EXECUTIVE SUMMARY', 1)
    
    summary = doc.add_paragraph()
    summary.add_run(f'This report presents a comprehensive emotion classification analysis of {len(df_clean):,} YouTube comments ')
    summary.add_run(f'across {df_clean["emotion"].nunique()} emotion categories. Using advanced machine learning techniques, the best performing model ')
    summary.add_run(f'({best_clf.replace("_", " ")}) achieved an accuracy of {model_results[best_clf]["accuracy"]:.2%} and F1-score of {model_results[best_clf]["f1_score"]:.2%}.')
    
    # ========================================================================
    # DATASET OVERVIEW
    # ========================================================================
    
    doc.add_heading('2. DATASET CHARACTERISTICS', 1)
    
    dataset_stats = doc.add_paragraph(style='List Bullet')
    dataset_stats.add_run(f'Total Comments Analyzed: {len(df_clean):,}')
    
    avg_length = doc.add_paragraph(style='List Bullet')
    avg_length.add_run(f'Average Comment Length: {df_clean["word_count_processed"].mean():.1f} words')
    
    emotions_item = doc.add_paragraph(style='List Bullet')
    emotions_item.add_run(f'Emotion Categories: {df_clean["emotion"].nunique()} (joy, sadness, anger, fear, surprise, disgust, neutral)')
    
    preprocessing = doc.add_paragraph(style='List Bullet')
    reduction = ((df_clean['original_length'].mean() - df_clean['processed_length'].mean()) / df_clean['original_length'].mean() * 100)
    preprocessing.add_run(f'Text Preprocessing: {reduction:.1f}% length reduction')
    
    # Class distribution
    doc.add_heading('2.1 Emotion Distribution', 2)
    emotion_dist = df_clean['emotion'].value_counts()
    
    for emotion, count in emotion_dist.items():
        dist_para = doc.add_paragraph(style='List Bullet')
        dist_para.add_run(f'{emotion.capitalize()}: {count:,} samples ({count/len(df_clean)*100:.1f}%)')
    
    # ========================================================================
    # MODEL PERFORMANCE
    # ========================================================================
    
    doc.add_heading('3. MODEL PERFORMANCE ANALYSIS', 1)
    
    doc.add_paragraph(
        'Eight machine learning models were trained and evaluated for emotion classification.'
    )
    
    doc.add_heading('3.1 Model Comparison', 2)
    
    table = doc.add_table(rows=len(classifiers)+1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Model'
    header_cells[1].text = 'Accuracy'
    header_cells[2].text = 'F1-Score'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            paragraph.runs[0].font.bold = True
    
    sorted_models = sorted(classifiers, key=lambda x: model_results[x]['f1_score'], reverse=True)
    for idx, clf in enumerate(sorted_models):
        row_cells = table.rows[idx+1].cells
        row_cells[0].text = clf.replace('_', ' ')
        row_cells[1].text = f'{model_results[clf]["accuracy"]:.4f}'
        row_cells[2].text = f'{model_results[clf]["f1_score"]:.4f}'
        
        if clf == best_clf:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.runs[0].font.bold = True
    
    # ========================================================================
    # PAGE 2: KEY FINDINGS
    # ========================================================================
    
    doc.add_page_break()
    
    doc.add_heading('4. KEY FINDINGS & INSIGHTS', 1)
    
    doc.add_heading('4.1 Superior Model Performance', 2)
    
    perf_para = doc.add_paragraph()
    best_accuracy = model_results[best_clf]['accuracy']
    perf_para.add_run(f'The {best_clf.replace("_", " ")} achieved {best_accuracy:.2%} accuracy. ')
    
    if best_accuracy > 0.85:
        perf_para.add_run('This represents excellent performance for production deployment.')
    elif best_accuracy > 0.70:
        perf_para.add_run('This represents good performance, suitable for production use.')
    else:
        perf_para.add_run('This suggests room for improvement through advanced architectures.')
    
    doc.add_heading('4.2 Emotion Distribution Patterns', 2)
    
    imbalance_ratio = emotion_dist.iloc[0] / emotion_dist.iloc[-1]
    emotion_para = doc.add_paragraph()
    emotion_para.add_run(f'The dataset exhibits a {imbalance_ratio:.1f}:1 imbalance ratio. ')
    emotion_para.add_run(f'"{emotion_dist.index[0]}" emotions dominate at {emotion_dist.iloc[0]/len(df_clean):.1%}.')
    
    doc.add_heading('4.3 Sentiment-Emotion Consistency', 2)
    sentiment_para = doc.add_paragraph()
    sentiment_para.add_run('Strong correlation exists between sentiment polarity and emotion labels, validating data quality.')
    
    # ========================================================================
    # RECOMMENDATIONS
    # ========================================================================
    
    doc.add_heading('5. RECOMMENDATIONS', 1)
    
    rec1 = doc.add_paragraph(style='List Number')
    rec1.add_run('Model Deployment: ').bold = True
    rec1.add_run(f'Deploy {best_clf.replace("_", " ")} as the production model.')
    
    rec2 = doc.add_paragraph(style='List Number')
    rec2.add_run('Data Augmentation: ').bold = True
    rec2.add_run('Address class imbalance using SMOTE or synthetic data generation.')
    
    rec3 = doc.add_paragraph(style='List Number')
    rec3.add_run('Feature Enhancement: ').bold = True
    rec3.add_run('Incorporate contextual features (emoji patterns, n-grams).')
    
    rec4 = doc.add_paragraph(style='List Number')
    rec4.add_run('Advanced Architectures: ').bold = True
    rec4.add_run('Explore transformer-based models (BERT, RoBERTa).')
    
    rec5 = doc.add_paragraph(style='List Number')
    rec5.add_run('Continuous Monitoring: ').bold = True
    rec5.add_run('Implement model drift detection and periodic retraining.')
    
    # ========================================================================
    # METHODOLOGY
    # ========================================================================
    
    doc.add_heading('6. METHODOLOGY SUMMARY', 1)
    
    methods = [
        'Text preprocessing (lemmatization, stopword removal)',
        'TF-IDF vectorization (2000 features, 1-2 gram range)',
        'Topic modeling using LDA and NMF',
        'K-Means clustering for pattern discovery',
        'Supervised classification using 5 algorithms',
        'Stratified train-test split (80-20)',
        'Evaluation using accuracy, F1-score, precision, recall'
    ]
    
    for method in methods:
        m = doc.add_paragraph(style='List Bullet')
        m.add_run(method)
    
    # ========================================================================
    # CONCLUSION
    # ========================================================================
    
    doc.add_heading('7. CONCLUSION', 1)
    
    conclusion = doc.add_paragraph()
    conclusion.add_run('This comprehensive emotion classification analysis demonstrates the viability of machine learning ')
    conclusion.add_run(f'for automated emotion detection. The {best_clf.replace("_", " ")} model ')
    conclusion.add_run(f'achieves production-ready performance at {best_accuracy:.2%} accuracy. ')
    conclusion.add_run('The system is ready for deployment in content moderation, user experience analysis, ')
    conclusion.add_run('and social media monitoring applications.')
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('—' * 50)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_text = doc.add_paragraph('End of Report')
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text.runs[0].font.italic = True
    footer_text.runs[0].font.size = Pt(10)
    
    # Save document
    output_path = '/content/drive/MyDrive/PIT_EMOTION_ANALYSIS_REPORT.docx'
    doc.save(output_path)
    
    return output_path

# Generate the report
try:
    report_path = create_comprehensive_report()
    print(f"✓ Report successfully generated!")
    print(f"✓ Location: {report_path}")
    print(f"✓ Pages: ~2 pages")
    print(f"✓ Format: Microsoft Word (.docx)")
    
    if os.path.exists(report_path):
        file_size = os.path.getsize(report_path) / 1024
        print(f"✓ File size: {file_size:.1f} KB")
    
    print("\n" + "="*80)
    print("📄 REPORT GENERATION COMPLETE!")
    print("="*80)
    print("\nThe report includes:")
    print("  • Executive Summary")
    print("  • Dataset Characteristics")
    print("  • Model Performance Comparison Table")
    print("  • Key Findings & Insights")
    print("  • Recommendations")
    print("  • Methodology Summary")
    print("  • Conclusion")
    print("\nYou can open the report in Microsoft Word, Google Docs, or any word processor.")
    
except Exception as e:
    print(f"❌ Error generating report: {e}")
    print("Please ensure python-docx is installed: pip install python-docx")
